import io
from docx import Document
import re

def clean_topic(text: str) -> str:
    """Mavzu matnidan oxiridagi nuqtalar, raqamlar va bo‘sh joylarni olib tashlaydi"""
    cleaned = re.sub(r"[.\s]*\.*\d+\s*$", "", text)
    return cleaned.strip()

def is_likely_topic(text: str) -> bool:
    """Matn katta harflarda yozilganmi yoki ilmiy mavzuga o‘xshaydimi"""
    return (
        text.isupper()
        or (len(text) > 20 and sum(c.isupper() for c in text) > 5)
    )

def is_likely_author(text: str) -> bool:
    """Muallif FIOsini aniqlash: lotin, kirill, rus, A.B.Ism va h.k."""
    patterns = [
        r"^[A-ZА-ЯЎҚҒҲ]\.[A-ZА-ЯЎҚҒҲ]\.\s?[A-ZА-ЯЎҚҒҲa-zа-ёўқғҳ’ʼʻ]+$",       # A.B. Ism
        r"^[A-ZА-ЯЎҚҒҲ][a-zа-ёўқғҳʼ’ʻ]+\s[A-ZА-ЯЎҚҒҲ][a-zа-ёўқғҳʼ’ʻ]+$",        # Ism Familiya
        r"^[A-ZА-ЯЎҚҒҲ][a-zа-ёўқғҳʼ’ʻ]+$",                                     # Faqat Ism yoki Familiya
    ]
    return any(re.match(p, text) for p in patterns)

def extract_all_text(doc: Document) -> list[str]:
    """Paragraf va jadval matnlarini birlashtirib beradi"""
    texts = []

    # Paragraflarni yig‘ish
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    # Jadval matnlarini ham qo‘shish
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        texts.append(text)

    return texts

def extract_fio_and_topics(file_bytes) -> list[tuple[str, str]]:
    doc = Document(io.BytesIO(file_bytes))
    texts = extract_all_text(doc)

    pairs = []
    current_authors = []

    for text in texts:
        if is_likely_author(text):
            current_authors.append(text)

        elif is_likely_topic(text):
            topic = clean_topic(text)
            for author in current_authors:
                pairs.append((author, topic))
            current_authors = []

        # Agar matn mavzuga ham, muallifga ham o‘xshamasa → tashlab ketiladi

    return pairs
